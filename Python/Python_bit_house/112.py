"""
FileName: 112.py
Author: Fatpandac
Create: 2021/11/20
Description: Reading a word file, return a paragraph with the most frequent 
             font style.
"""

from docx import Document

def main():
    doc = Document('data112.docx')
    doc_font_style = []
    for i in range(len(doc.paragraphs)):
        para_font_style = []
        for j in doc.paragraphs[i].runs:
            para_font_style.append(j.font.name)
        doc_font_style.append(para_font_style)
    doc_font_style = list(map(set, doc_font_style))
    return doc.paragraphs[doc_font_style.index(max(doc_font_style, key=len))].text

print(main())
